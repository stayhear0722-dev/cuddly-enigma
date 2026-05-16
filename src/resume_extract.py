from __future__ import annotations

from pathlib import Path


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def extract_resume_text(path: str | Path) -> str:
    resume_path = Path(path).expanduser().resolve()
    if not resume_path.exists():
        raise FileNotFoundError(f"Resume file not found: {resume_path}")
    if not resume_path.is_file():
        raise ValueError(f"Resume path is not a file: {resume_path}")

    suffix = resume_path.suffix.lower()
    if suffix == ".pdf":
        return normalize_text(extract_pdf_text(resume_path))
    if suffix == ".docx":
        return normalize_text(extract_docx_text(resume_path))
    if suffix == ".txt":
        return normalize_text(extract_txt_text(resume_path))
    raise ValueError(
        f"Unsupported resume format: {suffix}. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


def extract_pdf_text(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def extract_docx_text(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_txt_text(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gbk", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def normalize_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()
